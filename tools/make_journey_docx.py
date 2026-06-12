"""
Generate the leadership-ready Word document (Deloitte-branded) from the journey
content. Uses python-docx (no Node required).

Run:  C:\\Users\\surpanda\\tools\\python312\\python.exe tools\\make_journey_docx.py
Output: docs/Virtual-Engineer-Journey.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "docs" / "Virtual-Engineer-Journey.docx"

GREEN = RGBColor(0x86, 0xBC, 0x25)
GREEN_DARK = RGBColor(0x6F, 0x9E, 0x1F)
DARK = RGBColor(0x28, 0x27, 0x28)
GREY = RGBColor(0x5B, 0x64, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Open Sans"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK


def green_rule(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "86BC25")
    pbdr.append(bottom)
    pPr.append(pbdr)


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def heading(text, size=15, color=GREEN_DARK, before=16, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, italic=False, color=DARK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.italic = italic
    r.font.color.rgb = color
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.color.rgb = DARK
    return p


# --- Title block ---
top = doc.add_paragraph()
tr = top.add_run("Deloitte")
tr.bold = True
tr.font.size = Pt(12)
tr.font.color.rgb = DARK
dot = top.add_run(".")
dot.bold = True
dot.font.size = Pt(12)
dot.font.color.rgb = GREEN

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(0)
r1 = title.add_run("Virtual ")
r1.bold = True
r1.font.size = Pt(28)
r1.font.color.rgb = DARK
r2 = title.add_run("Engineer")
r2.bold = True
r2.italic = True
r2.font.size = Pt(28)
r2.font.color.rgb = GREEN_DARK

sub = doc.add_paragraph()
sr = sub.add_run("From idea to working product — a step-by-step account for the leadership team")
sr.font.size = Pt(12.5)
sr.font.color.rgb = GREY
green_rule(sub)
sub.paragraph_format.space_after = Pt(10)

date = doc.add_paragraph()
dr = date.add_run("June 2026  ·  Built with Claude Code")
dr.font.size = Pt(9.5)
dr.font.color.rgb = GREY

# --- Executive overview ---
heading("Executive overview")
body("In a matter of days, we built Virtual Engineer — an AI-powered incident-diagnosis copilot "
     "that helps L2/L3 engineers diagnose complex incidents in seconds instead of hours. It "
     "correlates incidents, changes, and problems across the ITSM estate and uses Claude to "
     "generate evidence-backed root-cause analysis, change-impact assessment, similar-incident "
     "guidance, and portfolio hotspots.")
body("The entire solution — code, data layer, branded interface, and documentation — was designed "
     "and built conversationally using Claude Code, on a standard locked-down corporate laptop, "
     "with no special infrastructure.")

# --- Steps ---
heading("How we got here, step by step")
steps = [
    ("Step 1 — Set up a working environment on a locked-down machine.",
     "Software installs are blocked by policy, so we used a portable, no-install toolchain.",
     "A fully working development setup with zero admin rights or exceptions."),
    ("Step 2 — Stood up source control.",
     "We versioned every change locally, then migrated the project to a public GitHub repository.",
     "Every step is tracked, reviewable, and shareable."),
    ("Step 3 — Built the first working app.",
     "A small, fast web application — the first Virtual Engineer — running entirely offline.",
     "A running product on day one, proving the approach."),
    ("Step 4 — Added document analysis.",
     "The ability to read PDFs, Word, Excel, PowerPoint, CSVs, and images and surface issues, requests, and trends.",
     "Turned unstructured files into structured insight."),
    ("Step 5 — Connected real AI (Claude).",
     "We wired in the Claude API, including reading images, with an automatic fallback so the app never breaks.",
     "Genuine, context-aware analysis — not canned responses."),
    ("Step 6 — Focused it on incident diagnosis (the core mission).",
     "We loaded ServiceNow ITSM data — tens of thousands of incidents, changes, problems, and tasks — "
     "into a fast, searchable database, correlated by the affected system.",
     "A single, queryable view across data that normally lives in separate screens."),
    ("Step 7 — Built four expert capabilities.",
     "Root Cause, Change Impact, Similar Incidents, and Hotspots & Trends — each pairing precise "
     "correlation with Claude's reasoning.",
     "Senior-engineer judgement, available to everyone, in seconds."),
    ("Step 8 — Made it modular and reusable.",
     "Pluggable data sources (a ready live ServiceNow connector) and a bring-your-own-data mode where "
     "uploaded files stay private to the session.",
     "Useful to any team, not just ours."),
    ("Step 9 — Made it professional and on-brand.",
     "Redesigned to the official Deloitte brand and consolidated onto one clean workspace, with grouped "
     "navigation, colour-coded charts, and result cards that each end with a clear Conclusion.",
     "An executive-ready, polished product."),
    ("Step 10 — Made it safe to share, and easy to demo.",
     "Added a synthetic demo dataset so anyone can run the app immediately, while real data and credentials "
     "stay private and are never published; plus a one-command demo mode and deployment-ready config.",
     "Shareable with confidence, with no data-governance risk."),
    ("Step 11 — Prepared leadership materials.",
     "A project summary, a recording-ready demo script with speaking notes, a one-minute talk track, and this document.",
     "Ready to present and socialise."),
]
for ttl, desc, out in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)
    rt = p.add_run(ttl)
    rt.bold = True
    rt.font.color.rgb = DARK
    body(desc, after=1)
    po = doc.add_paragraph()
    po.paragraph_format.space_after = Pt(4)
    ro = po.add_run("Outcome: " + out)
    ro.italic = True
    ro.font.color.rgb = GREEN_DARK

# --- At a glance table ---
heading("At a glance")
rows = [
    ("#", "Step", "What it delivered"),
    ("1", "Environment setup", "Working dev setup, no admin rights"),
    ("2", "Source control", "Tracked, shareable, public GitHub repo"),
    ("3", "First app", "Running product on day one"),
    ("4", "Document analysis", "Insight from unstructured files"),
    ("5", "Real Claude AI", "Genuine, context-aware analysis"),
    ("6", "ITSM data layer", "One correlated view across the estate"),
    ("7", "Four capabilities", "RCA, change impact, similar, hotspots"),
    ("8", "Modular + uploads", "Reusable by any team, private per session"),
    ("9", "Deloitte redesign", "Executive-ready, on-brand interface"),
    ("10", "Safe to share", "Demo data public; real data kept private"),
    ("11", "Leadership materials", "Summary, demo script, talk track"),
]
table = doc.add_table(rows=len(rows), cols=3)
table.style = "Table Grid"
widths = [Pt(28), Pt(150), Pt(270)]
for ri, row in enumerate(rows):
    cells = table.rows[ri].cells
    for ci, val in enumerate(row):
        cells[ci].width = widths[ci]
        para = cells[ci].paragraphs[0]
        run = para.add_run(val)
        run.font.size = Pt(10)
        if ri == 0:
            run.bold = True
            run.font.color.rgb = WHITE
            shade(cells[ci], "86BC25")
        else:
            run.font.color.rgb = DARK
            if ri % 2 == 0:
                shade(cells[ci], "F2F5EC")

# --- Value ---
heading("The value it produces")
for b in [
    "Speed — incident root cause goes from hours of manual cross-referencing to seconds, with the evidence cited.",
    "Scale — captures and reuses senior-engineer judgement, so every engineer operates at a higher level.",
    "Prevention — flags change-induced incidents before they spread.",
    "Reusability — bring-your-own-data for any team today; one configuration change from live ServiceNow and monitoring.",
    "Speed to build — delivered in days, not months, using Claude Code.",
]:
    bullet(b)

# --- What's next ---
heading("What's next")
for b in [
    "Connect live to ServiceNow and monitoring tools (metrics, logs, topology).",
    "Downloadable RCA and design reports.",
    "Optional secure hosting for broader, durable access.",
]:
    bullet(b)

close = doc.add_paragraph()
close.paragraph_format.space_before = Pt(12)
cr = close.add_run("Together makes progress — people and AI, diagnosing faster.")
cr.italic = True
cr.font.color.rgb = GREEN_DARK

repo = doc.add_paragraph()
rr = repo.add_run("Repository: github.com/surpanda1977/virtual-engineer")
rr.font.size = Pt(9.5)
rr.font.color.rgb = GREY

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print("Wrote", OUT)
